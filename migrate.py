"""
Migrate existing generated .docx files → PostgreSQL students table.
Run on the server: python3 migrate.py
"""
import os, re, uuid, glob
from datetime import datetime
from docx import Document
from docx.shared import Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from sqlalchemy import create_engine, Column, String, DateTime
from sqlalchemy.orm import declarative_base, sessionmaker
from dotenv import load_dotenv

load_dotenv()

GENERATED = os.path.join(os.path.dirname(__file__), 'generated')
UPLOADS   = os.path.join(os.path.dirname(__file__), 'uploads')
TEMPLATE  = os.path.join(os.path.dirname(__file__), 'imtihon_varaqasi (1).docx')

DATABASE_URL = os.getenv('DATABASE_URL', 'postgresql://localhost/stipendiya')
engine  = create_engine(DATABASE_URL)
Base    = declarative_base()
Session = sessionmaker(bind=engine)

class Student(Base):
    __tablename__ = 'students'
    id             = Column(String, primary_key=True)
    familiya       = Column(String, nullable=False)
    ism            = Column(String, nullable=False)
    otasining_ismi = Column(String, nullable=False)
    fakultet       = Column(String, nullable=False)
    yonalish       = Column(String, nullable=False)
    kurs           = Column(String, nullable=False)
    stipendiya     = Column(String, nullable=False)
    photo_path     = Column(String, nullable=True)
    doc_filename   = Column(String, nullable=True)
    created_at     = Column(DateTime, default=datetime.utcnow)

Base.metadata.create_all(engine)

# ── helpers ───────────────────────────────────────────────────────────────────
TUTUQ = re.compile(r"([oOgGaAuUeEiI])'")
_TUTUQ_CHAR = '\u2019'
def fix_q(t: str) -> str:
    return TUTUQ.sub(lambda m: m.group(1) + _TUTUQ_CHAR, t)

def strip_label(line: str) -> str:
    """'Fakultet:  Tabiiy fanlar' → 'Tabiiy fanlar'"""
    if ':' in line:
        return line.split(':', 1)[1].strip()
    return line.strip()

def extract_from_docx(path: str) -> dict:
    doc  = Document(path)
    info = doc.tables[1].rows[0].cells[0]
    paras = [p.text.strip() for p in info.paragraphs if p.text.strip()]

    # Para 0: "Ismi va familiya:  Familiya Ism Otasining_ismi"
    full_name = strip_label(paras[0]) if paras else ''
    parts = full_name.split()
    familiya       = parts[0] if len(parts) > 0 else ''
    ism            = parts[1] if len(parts) > 1 else ''
    otasining_ismi = ' '.join(parts[2:]) if len(parts) > 2 else ''

    # Para 1: Fakultet
    fakultet = strip_label(paras[1]) if len(paras) > 1 else ''

    # Para 2: Yo'nalish
    yonalish = strip_label(paras[2]) if len(paras) > 2 else ''

    # Para 3: Kurs
    kurs = ''
    if len(paras) > 3:
        kurs = strip_label(paras[3])
    if not kurs:
        kurs = '2-kurs'   # safe default

    # Stipendiya from header
    stipendiya = ''
    for para in doc.paragraphs:
        t = para.text
        if 'STIPENDIYA' in t and 'DAVLAT' in t:
            m = re.search(r'PREZIDENTI\s+(.+?)\s+DAVLAT', t, re.IGNORECASE)
            if m:
                stipendiya = m.group(1).strip().title()
            break

    return dict(
        familiya=familiya, ism=ism, otasining_ismi=otasining_ismi,
        fakultet=fakultet, yonalish=yonalish, kurs=kurs, stipendiya=stipendiya,
    )

def find_photo(uuid8: str) -> str | None:
    for f in os.listdir(UPLOADS):
        if f.startswith(uuid8):
            return os.path.join(UPLOADS, f)
    return None

# ── document re-generation ────────────────────────────────────────────────────
def replace_para(para, text, bold=None):
    first = para.runs[0] if para.runs else None
    for r in list(para.runs):
        r._element.getparent().remove(r._element)
    run = para.add_run(text)
    if bold is not None:
        run.bold = bold
    elif first:
        run.bold = first.bold
        if first.font.size: run.font.size = first.font.size

def insert_image(cell, path, w=3.0, h=4.0):
    for para in cell.paragraphs:
        for r in list(para.runs):
            r._element.getparent().remove(r._element)
    para = cell.paragraphs[0]
    pPr = para._element.find(qn('w:pPr'))
    if pPr is not None:
        sp = pPr.find(qn('w:spacing'))
        if sp is not None: pPr.remove(sp)
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        para._element.insert(0, pPr)
    sp_el = OxmlElement('w:spacing')
    sp_el.set(qn('w:before'), '0')
    sp_el.set(qn('w:after'),  '0')
    pPr.append(sp_el)
    para.alignment = 1
    para.add_run().add_picture(path, width=Cm(w), height=Cm(h))

def regenerate(s: Student, out_path: str):
    doc        = Document(TEMPLATE)
    info_table = doc.tables[1]
    info_cell  = info_table.rows[0].cells[0]
    photo_cell = info_table.rows[0].cells[1]

    full_name = fix_q(f"{s.familiya} {s.ism} {s.otasining_ismi}")
    paras = info_cell.paragraphs
    if len(paras) >= 1: replace_para(paras[0], f"Ismi va familiya:  {full_name}")
    if len(paras) >= 2: replace_para(paras[1], f"Fakultet:  {fix_q(s.fakultet)}")
    if len(paras) >= 3: replace_para(paras[2], f"Yo\u2019nalish:  {fix_q(s.yonalish)}")

    kp = OxmlElement('w:p')
    kr = OxmlElement('w:r')
    kt = OxmlElement('w:t')
    kt.text = f"Kurs:  {fix_q(s.kurs)}"
    kr.append(kt); kp.append(kr)
    info_cell._element.append(kp)

    if s.photo_path and os.path.exists(s.photo_path):
        insert_image(photo_cell, s.photo_path)

    stip = fix_q(s.stipendiya)
    for para in doc.paragraphs:
        if 'NOMLI' in para.text and 'STIPENDIYA' in para.text:
            replace_para(
                para,
                f"O\u2018ZBEKISTON RESPUBLIKA PREZIDENTI "
                f"{stip.upper()} DAVLAT STIPENDIYASI",
                bold=True,
            )
            break

    doc.save(out_path)

# ── MAIN ─────────────────────────────────────────────────────────────────────
def main():
    docs = sorted(glob.glob(os.path.join(GENERATED, '*.docx')))
    print(f"Found {len(docs)} documents to migrate.\n")

    with Session() as db:
        existing = {s.doc_filename for s in db.query(Student).all()}

    migrated = skipped = errors = 0

    for path in docs:
        fname = os.path.basename(path)
        if fname in existing:
            print(f"  SKIP (already in DB): {fname}")
            skipped += 1
            continue

        # UUID8 is the last underscore-separated part before .docx
        parts  = fname.replace('.docx', '').split('_')
        uuid8  = parts[-1]
        sid    = str(uuid.uuid4())

        # Override UUID so photo matching still works via uuid8
        photo = find_photo(uuid8)

        try:
            data = extract_from_docx(path)
        except Exception as e:
            print(f"  ERROR reading {fname}: {e}")
            errors += 1
            continue

        # Re-generate with fixed code
        new_fname = fname   # keep same filename
        try:
            s_tmp = type('S', (), {
                'familiya': data['familiya'], 'ism': data['ism'],
                'otasining_ismi': data['otasining_ismi'],
                'fakultet': data['fakultet'], 'yonalish': data['yonalish'],
                'kurs': data['kurs'], 'stipendiya': data['stipendiya'],
                'photo_path': photo,
            })()
            regenerate(s_tmp, path)  # overwrite in-place with fixed version
        except Exception as e:
            print(f"  WARNING regen failed {fname}: {e}")

        with Session() as db:
            s = Student(
                id             = sid,
                familiya       = data['familiya'],
                ism            = data['ism'],
                otasining_ismi = data['otasining_ismi'],
                fakultet       = data['fakultet'],
                yonalish       = data['yonalish'],
                kurs           = data['kurs'],
                stipendiya     = data['stipendiya'],
                photo_path     = photo,
                doc_filename   = new_fname,
                created_at     = datetime.utcnow(),
            )
            db.add(s); db.commit()

        print(f"  OK  {data['familiya']} {data['ism']} | {data['fakultet'][:30]} | {data['kurs']} | photo={'yes' if photo else 'NO'}")
        migrated += 1

    print(f"\n✓ Done: {migrated} migrated, {skipped} skipped, {errors} errors.")

if __name__ == '__main__':
    main()
