import os
import uuid
import json
import zipfile
import io
import shutil
import re
from copy import deepcopy
from flask import Flask, request, jsonify, send_file, render_template, send_from_directory
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import lxml.etree as etree

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['GENERATED_FOLDER'] = 'generated'
app.config['TEMPLATE'] = 'imtihon_varaqasi (1).docx'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024

# In-memory student list (survives while server is running)
students = []


def clean_uzbek_text(text):
    """Normalize Uzbek apostrophes/tutuq belgilari.
    If preceded by o, O, g, G -> replace with opening single quote ‘ (6 shape).
    Otherwise -> replace with closing single quote ’ (9 shape).
    """
    if not isinstance(text, str):
        return text
    apos_pattern = r"['‘’ʻʼ`´]"
    def repl(match):
        start = match.start()
        if start > 0 and text[start-1] in 'oOgG':
            return '‘'
        return '’'
    return re.sub(apos_pattern, repl, text)



def replace_paragraph_text(para, new_text, bold=None, font_size=None):
    """Replace all runs in a paragraph with a single run containing new_text."""
    # Keep formatting from first run
    first_run = para.runs[0] if para.runs else None

    # Clear all runs
    for run in para.runs:
        run._element.getparent().remove(run._element)

    # Add new run
    run = para.add_run(new_text)
    if bold is not None:
        run.bold = bold
    elif first_run:
        run.bold = first_run.bold
    if font_size:
        run.font.size = Pt(font_size)
    elif first_run and first_run.font.size:
        run.font.size = first_run.font.size
    return run


def insert_image_in_cell(cell, image_path, width_cm=3.0, height_cm=4.0):
    """Replace text in cell with an image, preserving table borders."""
    # Clear all runs from all paragraphs
    for para in cell.paragraphs:
        for run in list(para.runs):
            run._element.getparent().remove(run._element)

    # Remove extra spacing (before=700, after=700) that breaks layout
    para = cell.paragraphs[0]
    pPr = para._element.find(qn('w:pPr'))
    if pPr is not None:
        spacing = pPr.find(qn('w:spacing'))
        if spacing is not None:
            pPr.remove(spacing)
    # Add zero spacing explicitly
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        para._element.insert(0, pPr)
    spacing_el = OxmlElement('w:spacing')
    spacing_el.set(qn('w:before'), '0')
    spacing_el.set(qn('w:after'), '0')
    pPr.append(spacing_el)

    # Center the image
    para.alignment = 1  # WD_ALIGN_PARAGRAPH.CENTER

    run = para.add_run()
    run.add_picture(image_path, width=Cm(width_cm), height=Cm(height_cm))


def generate_document(student):
    """Generate a Word document for a student based on the template."""
    # Ensure text fields are cleaned
    familiya = clean_uzbek_text(student.get('familiya', ''))
    ism = clean_uzbek_text(student.get('ism', ''))
    otasining_ismi = clean_uzbek_text(student.get('otasining_ismi', ''))
    fakultet = clean_uzbek_text(student.get('fakultet', ''))
    yonalish = clean_uzbek_text(student.get('yonalish', ''))
    kurs = clean_uzbek_text(student.get('kurs', ''))
    stipendiya = clean_uzbek_text(student.get('stipendiya', ''))

    template_path = app.config['TEMPLATE']
    doc = Document(template_path)

    # --- Update student info table (table index 1) ---
    info_table = doc.tables[1]
    info_cell = info_table.rows[0].cells[0]
    photo_cell = info_table.rows[0].cells[1]

    full_name = f"{familiya} {ism} {otasining_ismi}"

    # Update paragraphs in info cell
    paragraphs = info_cell.paragraphs

    if len(paragraphs) >= 1:
        p = paragraphs[0]
        replace_paragraph_text(p, f"Ismi va familiya:  {full_name}")

    if len(paragraphs) >= 2:
        p = paragraphs[1]
        replace_paragraph_text(p, f"Fakultet:  {fakultet}")

    if len(paragraphs) >= 3:
        p = paragraphs[2]
        replace_paragraph_text(p, f"Yo\u2018nalish:  {yonalish}")

    # Add Kurs paragraph
    kurs_para = OxmlElement('w:p')
    kurs_r = OxmlElement('w:r')
    kurs_t = OxmlElement('w:t')
    kurs_t.text = f"Kurs:  {kurs}"
    kurs_r.append(kurs_t)
    kurs_para.append(kurs_r)
    info_cell._element.append(kurs_para)

    # --- Replace photo placeholder ---
    photo_path = student.get('photo_path')
    if photo_path and os.path.exists(photo_path):
        insert_image_in_cell(photo_cell, photo_path)

    # --- Update stipendiya name in paragraph 6 ---
    if stipendiya:
        for i, para in enumerate(doc.paragraphs):
            if 'NOMLI' in para.text and 'STIPENDIYA' in para.text:
                # Insert stipendiya name before "NOMLI DAVLAT STIPENDIYA"
                new_text = f"O\u2018ZBEKISTON RESPUBLIKASI PREZIDENTI {stipendiya.upper()} DAVLAT STIPENDIYASI"
                replace_paragraph_text(para, new_text, bold=True)
                break

    # Save to generated folder
    student_id = student['id']
    safe_name = f"{student['familiya']}_{student['ism']}".replace(' ', '_')
    filename = f"{safe_name}_{student_id[:8]}.docx"
    output_path = os.path.join(app.config['GENERATED_FOLDER'], filename)
    doc.save(output_path)

    return filename


@app.route('/')
def index():
    return render_template('index.html')


@app.route('/yuklab-olish')
def download_page():
    return render_template('download.html')


@app.route('/api/students', methods=['GET'])
def get_students():
    return jsonify(students)


@app.route('/api/students', methods=['POST'])
def add_student():
    data = request.form.to_dict()
    photo = request.files.get('photo')

    student_id = str(uuid.uuid4())
    photo_path = None

    if photo and photo.filename:
        ext = os.path.splitext(photo.filename)[1].lower()
        photo_filename = f"{student_id}{ext}"
        photo_path = os.path.join(app.config['UPLOAD_FOLDER'], photo_filename)
        photo.save(photo_path)

    student = {
        'id': student_id,
        'familiya': clean_uzbek_text(data.get('familiya', '').strip()),
        'ism': clean_uzbek_text(data.get('ism', '').strip()),
        'otasining_ismi': clean_uzbek_text(data.get('otasining_ismi', '').strip()),
        'fakultet': clean_uzbek_text(data.get('fakultet', '').strip()),
        'yonalish': clean_uzbek_text(data.get('yonalish', '').strip()),
        'kurs': clean_uzbek_text(data.get('kurs', '').strip()),
        'stipendiya': clean_uzbek_text(data.get('stipendiya', '').strip()),
        'photo_path': photo_path,
        'doc_filename': None,
    }

    # Generate document immediately
    try:
        filename = generate_document(student)
        student['doc_filename'] = filename
    except Exception as e:
        return jsonify({'error': str(e)}), 500

    students.append(student)

    # Return safe version (no internal path)
    safe = {k: v for k, v in student.items() if k != 'photo_path'}
    return jsonify(safe), 201


@app.route('/api/students/<student_id>', methods=['DELETE'])
def delete_student(student_id):
    global students
    student = next((s for s in students if s['id'] == student_id), None)
    if not student:
        return jsonify({'error': 'Not found'}), 404

    # Clean up files
    if student.get('photo_path') and os.path.exists(student['photo_path']):
        os.remove(student['photo_path'])
    if student.get('doc_filename'):
        doc_path = os.path.join(app.config['GENERATED_FOLDER'], student['doc_filename'])
        if os.path.exists(doc_path):
            os.remove(doc_path)

    students = [s for s in students if s['id'] != student_id]
    return jsonify({'ok': True})


@app.route('/api/download/<student_id>')
def download_student(student_id):
    student = next((s for s in students if s['id'] == student_id), None)
    if not student or not student.get('doc_filename'):
        return jsonify({'error': 'Not found'}), 404

    doc_path = os.path.join(app.config['GENERATED_FOLDER'], student['doc_filename'])
    if not os.path.exists(doc_path):
        return jsonify({'error': 'File not found'}), 404

    return send_file(
        doc_path,
        as_attachment=True,
        download_name=student['doc_filename'],
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )


@app.route('/api/download-all')
def download_all():
    if not students:
        return jsonify({'error': 'No students'}), 400

    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
        for student in students:
            if student.get('doc_filename'):
                doc_path = os.path.join(app.config['GENERATED_FOLDER'], student['doc_filename'])
                if os.path.exists(doc_path):
                    zf.write(doc_path, student['doc_filename'])

    zip_buffer.seek(0)
    return send_file(
        zip_buffer,
        as_attachment=True,
        download_name='imtihon_varaqalari.zip',
        mimetype='application/zip'
    )


@app.route('/api/regenerate/<student_id>', methods=['POST'])
def regenerate(student_id):
    student = next((s for s in students if s['id'] == student_id), None)
    if not student:
        return jsonify({'error': 'Not found'}), 404

    # Delete old doc
    if student.get('doc_filename'):
        old_path = os.path.join(app.config['GENERATED_FOLDER'], student['doc_filename'])
        if os.path.exists(old_path):
            os.remove(old_path)

    try:
        filename = generate_document(student)
        student['doc_filename'] = filename
    except Exception as e:
        return jsonify({'error': str(e)}), 500

    return jsonify({'ok': True, 'filename': filename})


if __name__ == '__main__':
    os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
    os.makedirs(app.config['GENERATED_FOLDER'], exist_ok=True)
    app.run(debug=True, port=5050)
